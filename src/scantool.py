import nmap
import re
import subprocess
import uuid
import os
import socket
import timeout_decorator
from selenium import webdriver
from selenium.webdriver.firefox.options import Options
import docx
from docx.shared import Inches
import requests

@timeout_decorator.timeout(60*60)
def _dirb(domain, wordlist):

    dirb_format = re.compile("^\+ (https?:\/\/.+) \(CODE:(\d+)\|SIZE:\d+\)$")
    dirb_format_dir = re.compile("^==> DIRECTORY: (https?://.+)$")
    retval = []
    random_name = str(uuid.uuid4())
    subprocess.run(["dirb", domain, "-o", random_name,"-w","-r"])
    with open(random_name,"r") as fp:
        for line in fp:
            match = dirb_format.match(line.strip())
            if match and int(match.group(2))!=301:
                retval.append(match.group(1))
            match = dirb_format_dir.match(line.strip())
            if match:
                retval.append(match.group(1))
    os.remove(random_name)

    return list(set(retval))

def dirb(domain, wordlist="/usr/share/dirb/wordlists/common.txt"):
    try:
        return _dirb(domain, wordlist)
    except timeout_decorator.TimeoutError:
        print("dirb timeout:",domain)
        return []


def nikto(url):
    random_name = os.path.join("tmp",str(uuid.uuid4())+".txt")
    subprocess.run(["nikto", "-h", url, "-o", random_name])
    return "\n".join(open(random_name,"r").readlines())


def get_screenshot(url, failure = "/assets/error.png"):
    for i in range(5):
        try:
            options = Options()
            options.headless = True
            driver = webdriver.Firefox(options=options)
            driver.get(url)
            driver.set_window_size(1920,1080)
            filename = re.sub("^https?:\/\/","",url)
            filename = filename.replace("/","-")+".png"
            driver.save_screenshot(filename)
            driver.quit()
            return filename
        except:
            continue
    return failure


def portscan(addr):
    print("port scannning...")
    ps = nmap.PortScanner()
    addr = socket.gethostbyname(addr)
    nmap_result = ps.scan(addr, arguments="-Pn -sS -T5 -A")
    print(ps.command_line())
    return nmap_result["scan"][addr]["tcp"]

def request_check(url_cand):
    try:
        requests.get(url_cand, timeout=10)
        return True
    except:
        return False



def portscan_str(nmap_result):
    retval=""
    for port in nmap_result:
        retval+=str(port)+"\n"
        for key in nmap_result[port]:
            retval+="\t{0}: {1}\n".format(key, nmap_result[port][key])
    return retval



def scan(addr, output="result"):
    doc = docx.Document()
    urls = []
    nmap_result = portscan(addr)
    doc.add_heading(addr, 0)
    doc.add_paragraph(portscan_str(nmap_result))
    doc.add_page_break()
    for port in nmap_result:
        if port!=443:
            url = "http://{0}:{1}".format(addr, port)
            if request_check(url):
                urls += dirb(url)
        if port!=80:
            url = "https://{0}:{1}".format(addr, port)
            if request_check(url):
                urls += dirb(url)
    screenshots = {url:get_screenshot(url) for url in urls}
    nikto_results = {url:nikto(url) for url in urls}


    for url in urls:
        doc.add_heading(url,0)
        doc.add_picture(screenshots[url], width=Inches(8.27*0.7))
        doc.add_heading("nikto",1)
        doc.add_paragraph(nikto_results[url])
        doc.add_page_break()
    doc.save("/mnt/{0}.docx".format(addr))

